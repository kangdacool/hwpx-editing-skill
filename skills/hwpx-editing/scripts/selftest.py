#!/usr/bin/env python3
"""
selftest.py — prove the core primitives work, no real HWPX needed.

Builds a tiny synthetic HWPX-shaped zip (mimetype STORED first, then deflated
XML entries) and checks that:
    1. a no-op repack is byte-identical to the source
    2. editing one entry re-deflates only that entry, keeps mimetype first/STORED,
       and leaves other entries byte-untouched
    3. adding a new entry (e.g. a BinData image) works and the zip stays valid
    4. duplicate-id detection and linesegarray stripping behave

Run:  python selftest.py    (exit 0 = all good)
"""
import io
import sys
import zipfile

# Force UTF-8 console output so non-ASCII glyphs in our messages (—, «», 한글)
# never crash on a Windows cp949 console (UnicodeEncodeError). No-op where the
# stream is already UTF-8, or where reconfigure() is unavailable (Python <3.7,
# or a redirected/replaced stream) — hence the broad guard.
try:
    sys.stdout.reconfigure(encoding="utf-8")
    sys.stderr.reconfigure(encoding="utf-8")
except (AttributeError, ValueError, OSError):
    pass

import hwpxlib as H


def build_fixture(path: str) -> None:
    zf = zipfile.ZipFile(path, "w")
    zi = zipfile.ZipInfo("mimetype")
    zi.compress_type = zipfile.ZIP_STORED
    zf.writestr(zi, b"application/hwp+zip")
    zf.writestr("Contents/header.xml",
                H.XML_DECL + b'<head xmlns="http://www.hancom.co.kr/hwpml/2011/head" itemCnt="3"/>')
    zf.writestr("Contents/section0.xml",
                H.XML_DECL + ('<sec xmlns:hp="http://www.hancom.co.kr/hwpml/2011/paragraph">'
                              '<hp:p id="10"><hp:linesegarray/>'
                              '<hp:run><hp:t>Hello 한글</hp:t></hp:run></hp:p>'
                              '<hp:p id="10"><hp:run><hp:t>dup id here</hp:t></hp:run></hp:p>'
                              '</sec>').encode("utf-8"))
    zf.writestr("Contents/content.hpf", H.XML_DECL + b"<package/>")
    zf.close()


def main() -> int:
    import tempfile
    import os
    d = tempfile.mkdtemp()
    orig = os.path.join(d, "orig.hwpx")
    build_fixture(orig)
    ok = True

    # 1. no-op byte-identity
    identical = H.self_verify_identical(orig)
    ok &= identical
    print(f"[{'PASS' if identical else 'FAIL'}] 1. no-op repack byte-identical")

    # 2. edit one entry
    new_sec = (H.XML_DECL +
               '<sec><hp:p xmlns:hp="http://www.hancom.co.kr/hwpml/2011/paragraph">'
               '<hp:run><hp:t>편집됨</hp:t></hp:run></hp:p></sec>'.encode("utf-8"))
    edited = os.path.join(d, "edited.hwpx")
    H.repack_preserve(orig, {"Contents/section0.xml": new_sec}, edited)
    z = zipfile.ZipFile(edited)
    zi = H.zip_integrity(z)
    e2 = (zi["testzip_ok"] and zi["mimetype_first"] and zi["mimetype_stored"]
          and z.read("Contents/section0.xml") == new_sec
          and z.read("Contents/header.xml") == zipfile.ZipFile(orig).read("Contents/header.xml"))
    ok &= e2
    print(f"[{'PASS' if e2 else 'FAIL'}] 2. edit re-deflates only the changed entry, "
          f"mimetype first/STORED, others untouched")

    # 3. add a new BinData entry
    img = b"\x89PNG_fake"
    added = os.path.join(d, "added.hwpx")
    H.repack_preserve(orig, {}, added, added={"BinData/image1.png": img})
    z3 = zipfile.ZipFile(added)
    e3 = z3.testzip() is None and z3.read("BinData/image1.png") == img
    ok &= e3
    print(f"[{'PASS' if e3 else 'FAIL'}] 3. add new entry (BinData) works, zip valid")

    # 4. duplicate-id detection + linesegarray strip
    root = H.etree.fromstring(zipfile.ZipFile(orig).read("Contents/section0.xml"))
    dups = H.find_duplicate_ids(root)
    removed = H.strip_linesegarray(root)
    e4 = (dups == {10: 2}) and (removed == 1)
    ok &= e4
    print(f"[{'PASS' if e4 else 'FAIL'}] 4. duplicate-id detection ({dups}) + "
          f"linesegarray strip ({removed} removed)")

    # 5. encoding regression (Windows cp949): the exact glyphs our CLIs print
    #    (— « » 한글) crash a raw cp949 stream, but the reconfigure(utf-8) guard
    #    the CLIs apply at import time must make them safe. Proves the fix works.
    glyphs = "— « » 한글"  # em-dash, guillemets, Hangul
    crashes_on_cp949 = False
    try:
        w = io.TextIOWrapper(io.BytesIO(), encoding="cp949")
        w.write(glyphs); w.flush()
    except UnicodeEncodeError:
        crashes_on_cp949 = True
    guard_ok = True
    try:
        w2 = io.TextIOWrapper(io.BytesIO(), encoding="cp949")
        try:
            w2.reconfigure(encoding="utf-8")
        except (AttributeError, ValueError, OSError):
            pass
        w2.write(glyphs); w2.flush()
    except UnicodeEncodeError:
        guard_ok = False
    e5 = crashes_on_cp949 and guard_ok
    ok &= e5
    print(f"[{'PASS' if e5 else 'FAIL'}] 5. cp949 console: raw stream crashes "
          f"({crashes_on_cp949}), reconfigure(utf-8) guard prints safely ({guard_ok})")

    # 6. own() must exclude 각주(footNote)/미주(endNote)/메모(fieldBegin) bodies,
    #    keeping only real body text.
    para_xml = (
        '<hp:p xmlns:hp="http://www.hancom.co.kr/hwpml/2011/paragraph">'
        '<hp:run><hp:t>본문앞</hp:t>'
        '<hp:ctrl><hp:footNote><hp:subList><hp:p><hp:run>'
        '<hp:t>각주내용</hp:t></hp:run></hp:p></hp:subList></hp:footNote></hp:ctrl>'
        '<hp:ctrl><hp:endNote><hp:subList><hp:p><hp:run>'
        '<hp:t>미주내용</hp:t></hp:run></hp:p></hp:subList></hp:endNote></hp:ctrl>'
        '<hp:fieldBegin><hp:subList><hp:p><hp:run>'
        '<hp:t>메모내용</hp:t></hp:run></hp:p></hp:subList></hp:fieldBegin>'
        '<hp:t>본문뒤</hp:t></hp:run></hp:p>'
    )
    body = H.own(H.etree.fromstring(para_xml.encode("utf-8")))
    e6 = (body == "본문앞본문뒤")
    ok &= e6
    print(f"[{'PASS' if e6 else 'FAIL'}] 6. own() excludes 각주/미주/메모 bodies "
          f"(got {body!r}, want '본문앞본문뒤')")

    # 7. duplicate-id policy: 한글 legitimately reuses an id on EMPTY structural
    #    paragraphs, so a pre-existing duplicate in the original is not the edit's
    #    fault. An edit that introduces NO new duplicate must not be flagged —
    #    only edit-introduced dupes (in edited but not in orig) are failures.
    pns = "http://www.hancom.co.kr/hwpml/2011/paragraph"
    orig_sec = H.etree.fromstring(
        (f'<sec xmlns:hp="{pns}">'
         '<hp:p id="7"><hp:run></hp:run></hp:p>'      # empty paragraph
         '<hp:p id="7"><hp:run></hp:run></hp:p>'      # same id reused on empty para
         '<hp:p id="9"><hp:run><hp:t>본문</hp:t></hp:run></hp:p>'
         '</sec>').encode("utf-8"))
    edited_sec = H.etree.fromstring(H.etree.tostring(orig_sec))
    H.etree.SubElement(edited_sec, f"{{{pns}}}p").set("id", "11")  # add unique id, no new dup
    orig_dups = set(H.find_duplicate_ids(orig_sec).keys())
    new_dups = {k: v for k, v in H.find_duplicate_ids(edited_sec).items()
                if k not in orig_dups}
    e7 = (orig_dups == {7}) and (new_dups == {})
    ok &= e7
    print(f"[{'PASS' if e7 else 'FAIL'}] 7. id-dup policy: orig dup {sorted(orig_dups)} "
          f"pre-existing, edit adds no new dup ({new_dups})")

    # 8. caption edit is surgical & lossless. HWPX table/figure captions are edited
    #    by replacing the last <hp:t> inside <hp:caption> (guide §4). Prove that doing
    #    so and repacking (a) updates only the caption text, (b) leaves other entries
    #    byte-identical, (c) yields a valid zip. (Rendering in 한글 is a manual check.)
    cap_sec = (H.XML_DECL + (
        f'<sec xmlns:hp="{pns}">'
        '<hp:p id="20"><hp:run><hp:ctrl><hp:tbl id="21"><hp:tr><hp:tc>'
        '<hp:subList><hp:p id="22"><hp:run><hp:t>셀</hp:t></hp:run></hp:p></hp:subList>'
        '</hp:tc></hp:tr>'
        '<hp:caption><hp:subList><hp:p id="23"><hp:run>'
        '<hp:t>표 1. 원래캡션</hp:t></hp:run></hp:p></hp:subList></hp:caption>'
        '</hp:tbl></hp:ctrl></hp:run></hp:p>'
        '</sec>').encode("utf-8"))
    cd = tempfile.mkdtemp()
    corig = os.path.join(cd, "cap.hwpx")
    zf = zipfile.ZipFile(corig, "w")
    zi = zipfile.ZipInfo("mimetype"); zi.compress_type = zipfile.ZIP_STORED
    zf.writestr(zi, b"application/hwp+zip")
    zf.writestr("Contents/header.xml",
                H.XML_DECL + b'<head xmlns="http://www.hancom.co.kr/hwpml/2011/head"/>')
    zf.writestr("Contents/section0.xml", cap_sec)
    zf.close()
    root = H.etree.fromstring(zipfile.ZipFile(corig).read("Contents/section0.xml"))
    cap_t = root.find(f".//{{{pns}}}caption").findall(f".//{{{pns}}}t")[-1]  # last <hp:t>
    cap_t.text = "표 1. 바뀐캡션"
    new_sec = H.XML_DECL + H.etree.tostring(root)
    cout = os.path.join(cd, "cap_edited.hwpx")
    H.repack_preserve(corig, {"Contents/section0.xml": new_sec}, cout)
    zc = zipfile.ZipFile(cout)
    got_cap = (H.etree.fromstring(zc.read("Contents/section0.xml"))
               .find(f".//{{{pns}}}caption").findall(f".//{{{pns}}}t")[-1].text)
    others_ok = zc.read("Contents/header.xml") == zipfile.ZipFile(corig).read("Contents/header.xml")
    diff = H.minimal_diff(cap_sec, new_sec)
    only_caption = bool(diff) and all("캡션" in ln for ln in diff)
    e8 = (got_cap == "표 1. 바뀐캡션") and others_ok and (zc.testzip() is None) and only_caption
    ok &= e8
    print(f"[{'PASS' if e8 else 'FAIL'}] 8. caption edit surgical+lossless "
          f"(cap now {got_cap!r}, {len(diff)} changed line(s), others byte-identical)")

    # 9. caption CREATION + position + alignment on an object that has none.
    #    Insert <hp:caption> right after <hp:outMargin> (ShapeObject order
    #    sz·pos·outMargin·caption·…); position via side=, text alignment via the
    #    caption paragraph's paraPrIDRef. Prove it lands in the right slot, carries
    #    the chosen side/alignment/text, and is lossless. (Render = manual check.)
    tbl_sec = (H.XML_DECL + (
        f'<sec xmlns:hp="{pns}">'
        '<hp:p id="30"><hp:run><hp:ctrl><hp:tbl id="31">'
        '<hp:sz width="40000" height="2000"/><hp:pos/><hp:outMargin/>'
        '<hp:tr><hp:tc><hp:subList><hp:p id="0"><hp:run charPrIDRef="1">'
        '<hp:t>셀</hp:t></hp:run></hp:p></hp:subList></hp:tc></hp:tr>'
        '</hp:tbl></hp:ctrl></hp:run></hp:p></sec>').encode("utf-8"))
    cd2 = tempfile.mkdtemp()
    torig = os.path.join(cd2, "t.hwpx")
    zf = zipfile.ZipFile(torig, "w")
    zi = zipfile.ZipInfo("mimetype"); zi.compress_type = zipfile.ZIP_STORED
    zf.writestr(zi, b"application/hwp+zip")
    zf.writestr("Contents/header.xml",
                H.XML_DECL + b'<head xmlns="http://www.hancom.co.kr/hwpml/2011/head"/>')
    zf.writestr("Contents/section0.xml", tbl_sec); zf.close()
    r9 = H.etree.fromstring(zipfile.ZipFile(torig).read("Contents/section0.xml"))
    tbl9 = r9.find(f".//{{{pns}}}tbl")
    cap = H.etree.Element(f"{{{pns}}}caption"); cap.set("side", "TOP")  # position
    sub9 = H.etree.SubElement(cap, f"{{{pns}}}subList")
    cp9 = H.etree.SubElement(sub9, f"{{{pns}}}p"); cp9.set("paraPrIDRef", "1")  # alignment
    crun9 = H.etree.SubElement(cp9, f"{{{pns}}}run"); crun9.set("charPrIDRef", "1")
    H.etree.SubElement(crun9, f"{{{pns}}}t").text = "표 1. 생성된 캡션"
    idx = next(i for i, c in enumerate(list(tbl9))
               if H.etree.QName(c).localname == "outMargin")
    tbl9.insert(idx + 1, cap)  # right after outMargin
    tnew = H.XML_DECL + H.etree.tostring(r9)
    tout = os.path.join(cd2, "t_edited.hwpx")
    H.repack_preserve(torig, {"Contents/section0.xml": tnew}, tout)
    zt = zipfile.ZipFile(tout)
    tb = H.etree.fromstring(zt.read("Contents/section0.xml")).find(f".//{{{pns}}}tbl")
    order = [H.etree.QName(c).localname for c in tb]
    capn = tb.find(f"{{{pns}}}caption")
    align9 = capn.find(f".//{{{pns}}}p").get("paraPrIDRef") if capn is not None else None
    e9 = (order[:4] == ["sz", "pos", "outMargin", "caption"]
          and capn is not None and capn.get("side") == "TOP" and align9 == "1"
          and capn.find(f".//{{{pns}}}t").text == "표 1. 생성된 캡션"
          and zt.testzip() is None
          and zt.read("Contents/header.xml") == zipfile.ZipFile(torig).read("Contents/header.xml"))
    ok &= e9
    print(f"[{'PASS' if e9 else 'FAIL'}] 9. caption create+position+align "
          f"(order={order[:4]}, side={capn.get('side') if capn is not None else None}, align={align9})")

    # 10. table → Excel export preserves merged cells. Build a table with BOTH a
    #     colSpan (title row) and a rowSpan (side header), convert via
    #     tables_to_xlsx, and confirm the .xlsx carries both merges + text.
    #     Skipped (not failed) when openpyxl is absent — it's an optional export dep.
    try:
        import openpyxl  # noqa: F401
        import tables_to_xlsx as TX
        _have_xlsx = True
    except ImportError:
        _have_xlsx = False
    if _have_xlsx:
        def _tc(col, row, cs, rs, txt):
            return (f'<hp:tc><hp:cellAddr colAddr="{col}" rowAddr="{row}"/>'
                    f'<hp:cellSpan colSpan="{cs}" rowSpan="{rs}"/><hp:subList><hp:p>'
                    f'<hp:run charPrIDRef="0"><hp:t>{txt}</hp:t></hp:run></hp:p></hp:subList></hp:tc>')
        tsec = (H.XML_DECL + (
            f'<sec xmlns:hp="{pns}"><hp:p id="40"><hp:run><hp:ctrl>'
            '<hp:tbl id="41" rowCnt="3" colCnt="3">'
            '<hp:sz width="30000" height="3000"/><hp:pos/><hp:outMargin/>'
            '<hp:tr>' + _tc(0, 0, 3, 1, "제목") + '</hp:tr>'
            '<hp:tr>' + _tc(0, 1, 1, 2, "세로") + _tc(1, 1, 1, 1, "b") + _tc(2, 1, 1, 1, "c") + '</hp:tr>'
            '<hp:tr>' + _tc(1, 2, 1, 1, "e") + _tc(2, 2, 1, 1, "f") + '</hp:tr>'
            '</hp:tbl></hp:ctrl></hp:run></hp:p></sec>').encode("utf-8"))
        cd3 = tempfile.mkdtemp()
        thwpx = os.path.join(cd3, "t.hwpx")
        zf = zipfile.ZipFile(thwpx, "w")
        zi = zipfile.ZipInfo("mimetype"); zi.compress_type = zipfile.ZIP_STORED
        zf.writestr(zi, b"application/hwp+zip")
        zf.writestr("Contents/header.xml",
                    H.XML_DECL + b'<head xmlns="http://www.hancom.co.kr/hwpml/2011/head"/>')
        zf.writestr("Contents/section0.xml", tsec); zf.close()
        xout = os.path.join(cd3, "t.xlsx")
        TX.convert(thwpx, xout)
        ws = openpyxl.load_workbook(xout)["T1"]
        ranges = {str(m) for m in ws.merged_cells.ranges}
        e10 = ("A1:C1" in ranges and "A2:A3" in ranges
               and ws["A1"].value == "제목" and ws["A2"].value == "세로" and ws["C3"].value == "f")
        ok &= e10
        print(f"[{'PASS' if e10 else 'FAIL'}] 10. table→xlsx merges preserved "
              f"(ranges={sorted(ranges)}, A1={ws['A1'].value!r})")
    else:
        print("[SKIP] 10. table→xlsx (openpyxl not installed — pip install openpyxl)")

    # 11. IDRef/itemCnt integrity: a dangling charPrIDRef or a stale itemCnt must be
    #     caught; a clean file must pass (validated against real 한글 files — no
    #     false positives).
    hns2 = "http://www.hancom.co.kr/hwpml/2011/head"
    def _idref_zip(path, itemcnt, ref):
        zf = zipfile.ZipFile(path, "w")
        zi = zipfile.ZipInfo("mimetype"); zi.compress_type = zipfile.ZIP_STORED
        zf.writestr(zi, b"application/hwp+zip")
        zf.writestr("Contents/header.xml", H.XML_DECL + (
            f'<hh:head xmlns:hh="{hns2}"><hh:refList>'
            f'<hh:charProperties itemCnt="{itemcnt}"><hh:charPr id="0"/></hh:charProperties>'
            f'</hh:refList></hh:head>').encode("utf-8"))
        zf.writestr("Contents/section0.xml", H.XML_DECL + (
            f'<sec xmlns:hp="{pns}"><hp:p><hp:run charPrIDRef="{ref}">'
            f'<hp:t>x</hp:t></hp:run></hp:p></sec>').encode("utf-8"))
        zf.close()
    dd = tempfile.mkdtemp()
    bad = os.path.join(dd, "bad.hwpx"); _idref_zip(bad, "2", "9")    # itemCnt 2 vs 1 child; ref 9 dangling
    good = os.path.join(dd, "good.hwpx"); _idref_zip(good, "1", "0")  # itemCnt ok; ref 0 exists
    rb = H.check_idref_integrity(zipfile.ZipFile(bad))
    rg = H.check_idref_integrity(zipfile.ZipFile(good))
    e11 = (len(rb["itemcnt"]) == 1 and len(rb["dangling"]) == 1
           and rg["itemcnt"] == [] and rg["dangling"] == [])
    ok &= e11
    print(f"[{'PASS' if e11 else 'FAIL'}] 11. IDRef/itemCnt integrity: bad flagged "
          f"(itemcnt={len(rb['itemcnt'])} dangling={len(rb['dangling'])}), clean passes")

    # 12. hwpx → Markdown extraction: body text, an inline footnote, and a table
    #     must all render (lxml only — no extra deps).
    import hwpx_to_markdown as MD
    def _c(col, row, txt):
        return (f'<hp:tc><hp:cellAddr colAddr="{col}" rowAddr="{row}"/>'
                f'<hp:cellSpan colSpan="1" rowSpan="1"/><hp:subList><hp:p><hp:run>'
                f'<hp:t>{txt}</hp:t></hp:run></hp:p></hp:subList></hp:tc>')
    mdsec = (H.XML_DECL + (
        f'<sec xmlns:hp="{pns}">'
        '<hp:p><hp:run><hp:t>본문 텍스트</hp:t>'
        '<hp:ctrl><hp:footNote><hp:subList><hp:p><hp:run><hp:t>각주내용</hp:t>'
        '</hp:run></hp:p></hp:subList></hp:footNote></hp:ctrl></hp:run></hp:p>'
        '<hp:p><hp:run><hp:ctrl><hp:tbl rowCnt="2" colCnt="2">'
        '<hp:tr>' + _c(0, 0, "A") + _c(1, 0, "B") + '</hp:tr>'
        '<hp:tr>' + _c(0, 1, "C") + _c(1, 1, "D") + '</hp:tr>'
        '</hp:tbl></hp:ctrl></hp:run></hp:p></sec>').encode("utf-8"))
    dd2 = tempfile.mkdtemp(); mpath = os.path.join(dd2, "m.hwpx")
    zf = zipfile.ZipFile(mpath, "w")
    zi = zipfile.ZipInfo("mimetype"); zi.compress_type = zipfile.ZIP_STORED
    zf.writestr(zi, b"application/hwp+zip")
    zf.writestr("Contents/section0.xml", mdsec); zf.close()
    md = MD.to_markdown(mpath)
    e12 = ("본문 텍스트" in md and "(각주: 각주내용)" in md
           and "| A | B |" in md and "| --- | --- |" in md and "| C | D |" in md)
    ok &= e12
    print(f"[{'PASS' if e12 else 'FAIL'}] 12. hwpx→markdown (body + inline footnote + table): "
          f"{'ok' if e12 else repr(md[:90])}")

    # 13. Excel/CSV → HWPX table insert: a rowSpan merge from data must produce a
    #     cell with the right span AND summed height, cloning the target's table,
    #     and omit the covered cell. (lxml only — insert_table needs no openpyxl.)
    import data_to_hwpx_table as DT
    tgt_sec = (H.XML_DECL + (
        f'<sec xmlns:hp="{pns}"><hp:p><hp:run><hp:ctrl>'
        '<hp:tbl id="50" rowCnt="1" colCnt="1"><hp:sz width="40000" height="300"/>'
        '<hp:pos/><hp:outMargin/><hp:inMargin/><hp:tr><hp:tc>'
        '<hp:subList><hp:p id="0"><hp:run charPrIDRef="0"><hp:t>x</hp:t></hp:run></hp:p></hp:subList>'
        '<hp:cellAddr colAddr="0" rowAddr="0"/><hp:cellSpan colSpan="1" rowSpan="1"/>'
        '<hp:cellSz width="40000" height="300"/>'
        '<hp:cellMargin left="0" right="0" top="0" bottom="0"/></hp:tc></hp:tr>'
        '</hp:tbl></hp:ctrl></hp:run></hp:p></sec>').encode("utf-8"))
    cd4 = tempfile.mkdtemp(); tgt = os.path.join(cd4, "t.hwpx")
    zf = zipfile.ZipFile(tgt, "w")
    zi = zipfile.ZipInfo("mimetype"); zi.compress_type = zipfile.ZIP_STORED
    zf.writestr(zi, b"application/hwp+zip")
    zf.writestr("Contents/section0.xml", tgt_sec); zf.close()
    tout = os.path.join(cd4, "t_out.hwpx")
    DT.insert_table(tgt, [["병합", "b"], ["", "d"]], [(0, 0, 2, 1)], tout)
    q = f"{{{pns}}}"
    nt = H.etree.fromstring(zipfile.ZipFile(tout).read("Contents/section0.xml")).findall(f".//{q}tbl")[-1]
    c0 = nt.find(f".//{q}tc")  # first cell = (0,0), the rowSpan origin
    e13 = (nt.get("rowCnt") == "2" and nt.get("colCnt") == "2"
           and c0.find(f"{q}cellSpan").get("rowSpan") == "2"
           and c0.find(f"{q}cellSz").get("height") == "600"          # 2 * 300 summed
           and "".join(c0.itertext()) == "병합"
           and len(nt.findall(f"{q}tr")[1].findall(f"{q}tc")) == 1)   # row1: covered col0 omitted
    ok &= e13
    print(f"[{'PASS' if e13 else 'FAIL'}] 13. data→hwpx table insert "
          f"(rowSpan span+summed height, covered cell omitted)")

    # 14. hwpx → Word (.docx): Hancom private-use glyphs stripped, and adjacent
    #     tables kept separate (Word merges back-to-back tables). strip_pua is core
    #     (lxml); the .docx build is skipped without python-docx.
    assert H.strip_pua("ab") == "ab"  # PUA (category Co) removed
    try:
        import docx  # noqa: F401
        import hwpx_to_docx as DX
        _have_docx = True
    except ImportError:
        _have_docx = False
    if _have_docx:
        dsec = (H.XML_DECL + (
            f'<sec xmlns:hp="{pns}">'
            '<hp:p><hp:run><hp:t>제목</hp:t></hp:run></hp:p>'
            '<hp:p><hp:run><hp:ctrl><hp:tbl rowCnt="1" colCnt="2"><hp:tr>'
            + _c(0, 0, "A") + _c(1, 0, "B") + '</hp:tr></hp:tbl></hp:ctrl></hp:run></hp:p>'
            '<hp:p><hp:run><hp:ctrl><hp:tbl rowCnt="1" colCnt="2"><hp:tr>'
            + _c(0, 0, "C") + _c(1, 0, "D") + '</hp:tr></hp:tbl></hp:ctrl></hp:run></hp:p>'
            '</sec>').encode("utf-8"))
        cd5 = tempfile.mkdtemp(); dhwpx = os.path.join(cd5, "d.hwpx")
        zf = zipfile.ZipFile(dhwpx, "w")
        zi = zipfile.ZipInfo("mimetype"); zi.compress_type = zipfile.ZIP_STORED
        zf.writestr(zi, b"application/hwp+zip")
        zf.writestr("Contents/section0.xml", dsec); zf.close()
        DX.convert(dhwpx, os.path.join(cd5, "d.docx"))
        d = docx.Document(os.path.join(cd5, "d.docx"))
        texts = [p.text for p in d.paragraphs]
        kids = [c.tag.split("}")[-1] for c in d.element.body]
        adj = sum(1 for i in range(len(kids) - 1) if kids[i] == "tbl" and kids[i + 1] == "tbl")
        e14 = ("제목" in texts and "제목" not in texts
               and len(d.tables) == 2 and adj == 0)
        ok &= e14
        print(f"[{'PASS' if e14 else 'FAIL'}] 14. hwpx→docx "
              f"(PUA stripped, {len(d.tables)} tables, adjacent tbl→tbl={adj})")
    else:
        print("[SKIP] 14. hwpx→docx (python-docx not installed — pip install python-docx)")

    print()
    print("RESULT:", "ALL PASS" if ok else "FAILURES PRESENT")
    return 0 if ok else 1


if __name__ == "__main__":
    sys.exit(main())
