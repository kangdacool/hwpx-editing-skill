#!/usr/bin/env python3
"""
hwpx_to_docx.py — export an HWPX's content to a Word document (.docx), so you can
hand it to a journal / co-authors who don't use 한글, or edit it in Word.

Usage:
    python hwpx_to_docx.py FILE.hwpx [OUT.docx]   # default: <file>.docx

Body paragraphs become Word paragraphs; tables become Word tables with merged
cells preserved (colSpan/rowSpan). A 1x1 layout frame is written as plain text;
footnotes/endnotes are appended inline as (각주: …) / (미주: …); an image/equation
-only cell shows [그림] / [수식]. Hancom private-use glyphs (custom bullets that
render as broken boxes outside 한글) are stripped.

Requires: lxml + python-docx  (pip install python-docx)
"""
import argparse
import sys
import zipfile

try:
    sys.stdout.reconfigure(encoding="utf-8")
    sys.stderr.reconfigure(encoding="utf-8")
except (AttributeError, ValueError, OSError):
    pass

import hwpxlib as H

try:
    import docx
except ImportError as e:  # pragma: no cover
    raise SystemExit("hwpx_to_docx needs python-docx:  pip install python-docx") from e

P = H.P
_SKIP = (f"{P}tc", f"{P}footNote", f"{P}endNote", f"{P}fieldBegin")


def _para_text(p) -> str:
    return H.strip_pua("".join("".join(t.itertext()) for t in p.findall(f".//{P}t")
                       if not any(a.tag in _SKIP for a in t.iterancestors())).strip())


def _cell_text(tc) -> str:
    parts = ["".join(t.itertext()) for t in tc.findall(f".//{P}t")
             if next((a for a in t.iterancestors() if a.tag == f"{P}tc"), None) is tc]
    txt = H.strip_pua(" ".join(" ".join(parts).split()))
    if not txt and tc.find(f".//{P}equation") is not None:
        txt = "[수식]"
    if not txt and tc.find(f".//{P}pic") is not None:
        txt = "[그림]"
    return txt


def _notes(p) -> str:
    out = ""
    for tag, label in ((f"{P}footNote", "각주"), (f"{P}endNote", "미주")):
        for note in p.findall(f".//{tag}"):
            txt = H.strip_pua(" ".join("".join(note.itertext()).split()))
            if txt:
                out += f" ({label}: {txt})"
    return out


def _add_table(doc, tbl):
    rc = int(tbl.get("rowCnt") or 0)
    cc = int(tbl.get("colCnt") or 0)
    if rc < 1 or cc < 1:
        return
    t = doc.add_table(rows=rc, cols=cc)
    t.style = "Table Grid"
    for tc in tbl.findall(f".//{P}tc"):
        if next((a for a in tc.iterancestors() if a.tag == f"{P}tc"), None) is not None:
            continue
        addr = tc.find(f"{P}cellAddr")
        span = tc.find(f"{P}cellSpan")
        r, c = int(addr.get("rowAddr")), int(addr.get("colAddr"))
        rs = int((span.get("rowSpan") if span is not None else "1") or 1)
        cs = int((span.get("colSpan") if span is not None else "1") or 1)
        if r >= rc or c >= cc:
            continue
        cell = t.cell(r, c)
        if rs > 1 or cs > 1:
            cell = cell.merge(t.cell(min(r + rs - 1, rc - 1), min(c + cs - 1, cc - 1)))
        cell.text = _cell_text(tc)


def convert(path, out):
    """Write the HWPX body to `out` (.docx). Returns (paragraph_count, table_count)."""
    H.ensure_hwpx(path)
    z = zipfile.ZipFile(path)
    doc = docx.Document()
    last_was_table = False
    for name in H.section_names(z):
        root = H.etree.fromstring(z.read(name))
        for p in root.findall(f"{P}p"):
            text = (_para_text(p) + _notes(p)).strip()
            if text:
                doc.add_paragraph(text)
                last_was_table = False
            for tbl in p.findall(f".//{P}tbl"):
                if next((a for a in tbl.iterancestors() if a.tag == f"{P}tbl"), None) is not None:
                    continue
                rc = int(tbl.get("rowCnt") or 0)
                cc = int(tbl.get("colCnt") or 0)
                if rc <= 1 and cc <= 1:
                    one = tbl.find(f".//{P}tc")
                    cell = _cell_text(one) if one is not None else ""
                    if cell and not text:
                        doc.add_paragraph(cell)
                        last_was_table = False
                else:
                    if last_was_table:
                        doc.add_paragraph("")  # separate adjacent tables (Word merges them otherwise)
                    _add_table(doc, tbl)
                    last_was_table = True
    doc.save(out)
    return len(doc.paragraphs), len(doc.tables)


def main() -> int:
    import os
    ap = argparse.ArgumentParser(description="Export HWPX content to Word (.docx).")
    ap.add_argument("file", help="the .hwpx to read")
    ap.add_argument("out", nargs="?", help="output .docx (default: <file>.docx)")
    args = ap.parse_args()
    out = args.out or (os.path.splitext(args.file)[0] + ".docx")
    try:
        n, tbls = convert(args.file, out)
    except H.NotHwpxError as e:
        print(str(e))
        return 2
    except FileNotFoundError:
        print(f"File not found: {args.file}")
        return 2
    print(f"Wrote {out}  ({n} paragraphs, {tbls} tables)")
    return 0


if __name__ == "__main__":
    sys.exit(main())
